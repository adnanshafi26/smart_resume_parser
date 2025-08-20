from docx import Document

# Create a new Document
doc = Document()

# Add Name and Contact
doc.add_heading("John Doe", 0)
doc.add_paragraph("Email: johndoe@example.com | Phone: +91-9876543210 | LinkedIn: linkedin.com/in/johndoe")

# Add Summary
doc.add_heading("Summary", level=1)
doc.add_paragraph(
    "Enthusiastic Software Engineer with 3+ years of experience in full-stack development. "
    "Skilled in Python, JavaScript, and cloud technologies. Passionate about building scalable applications."
)

# Add Education
doc.add_heading("Education", level=1)
doc.add_paragraph("B.Tech in Computer Science, XYZ University, 2018 - 2022\nCGPA: 8.5/10")

# Add Experience
doc.add_heading("Experience", level=1)
doc.add_paragraph(
    "Software Engineer, ABC Technologies (2022 - Present)\n"
    "- Developed REST APIs in Python Flask\n"
    "- Built interactive dashboards with React.js\n"
    "- Collaborated with a 5-member team in Agile environment"
)

doc.add_paragraph(
    "Intern, DEF Corp (Jan 2022 - Jun 2022)\n"
    "- Automated data processing using Pandas\n"
    "- Assisted in migrating on-prem systems to AWS"
)

# Add Skills
doc.add_heading("Skills", level=1)
doc.add_paragraph("Python, JavaScript, Flask, React.js, SQL, MongoDB, AWS, Git")

# Save the file
doc.save("sample_resume.docx")

print("✅ Sample resume created: sample_resume.docx")
